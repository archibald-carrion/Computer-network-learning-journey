from scapy.all import Ether, IP, TCP, Raw, wrpcap
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- PART 1: GENERATE THE PCAP ---
def create_tcp_session(c_ip, s_ip, c_port, s_port, seq_c, seq_s, data_list):
    sess = []
    # Handshake
    sess.append(Ether()/IP(src=c_ip, dst=s_ip)/TCP(sport=c_port, dport=s_port, flags="S", seq=seq_c))
    sess.append(Ether()/IP(src=s_ip, dst=c_ip)/TCP(sport=s_port, dport=c_port, flags="SA", seq=seq_s, ack=seq_c+1))
    sess.append(Ether()/IP(src=c_ip, dst=s_ip)/TCP(sport=c_port, dport=s_port, flags="A", seq=seq_c+1, ack=seq_s+1))
    
    curr_c = seq_c + 1
    curr_s = seq_s + 1
    
    # Data Exchange
    for data in data_list:
        p = Ether()/IP(src=s_ip, dst=c_ip)/TCP(sport=s_port, dport=c_port, flags="PA", seq=curr_s, ack=curr_c)/Raw(load=data)
        sess.append(p)
        curr_s += len(data)
    return sess

packets = []
# Chat Flow: Small packets on Port 80
packets += create_tcp_session("10.0.0.1", "10.0.0.2", 4444, 80, 1000, 5000, ["Hi", "Requesting tiny file", "Done"])

# Video Flow: Large MTU-sized packets on Port 8080
video_chunk = "V" * 1460 
packets += create_tcp_session("10.0.0.1", "10.0.0.2", 5555, 8080, 2000, 9000, [video_chunk] * 30)

wrpcap("mixed_traffic.pcap", packets)

# --- PART 2: GENERATE THE DOCX ---
doc = Document()
title = doc.add_heading('Assignment: Traffic Characterization & Flow Analysis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Objective', level=1)
doc.add_paragraph('Distinguish between application types by analyzing flow statistics and packet size distribution in a mixed-traffic environment.')

sections = {
    "1. The 'Conversations' Snapshot": [
        "Open mixed_traffic.pcap and navigate to Statistics -> Conversations -> TCP.",
        "Compare the 'Bytes' and 'Packets' columns for both rows. Which Stream ID corresponds to the 'Chat' and which to the 'Video'? Justify your answer using the data.",
        "Calculate the average packet size for each stream. How does this relate to the application's likely purpose?"
    ],
    "2. Packet Size Distribution": [
        "Navigate to Statistics -> Packet Lengths.",
        "Identify the two distinct spikes in the distribution graph. Explain what each represents in the context of these two applications.",
        "Why does the 'Chat' application result in significantly more small (control) overhead compared to the 'Video' stream?"
    ],
    "3. Throughput & I/O Graphs": [
        "Open Statistics -> I/O Graphs. Create filters for 'tcp.port == 80' and 'tcp.port == 8080'.",
        "Set the Y-Axis to 'Bits/Sec'. What is the peak throughput of each stream?",
        "If you were to apply a Machine Learning model to classify these flows, which three statistical features would be most reliable for identification?"
    ]
}

for section_title, questions in sections.items():
    doc.add_heading(section_title, level=2)
    for q in questions:
        doc.add_paragraph(q, style='List Number')

doc.save('Mixed_Traffic_Assignment.docx')
print("Success: 'mixed_traffic.pcap' and 'Mixed_Traffic_Assignment.docx' have been created.")
